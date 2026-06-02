"""Tests for document extraction API endpoint /files/extract-document.

Covers:
- .docx extraction (Word → Markdown)
- .xlsx extraction (Excel → Markdown table)
- Validation errors (unsupported format, missing params)
- filePath mode for local files
"""

import os

import pytest
from httpx import ASGITransport, AsyncClient

from app.main import app


@pytest.fixture
def anyio_backend():
    return "asyncio"


@pytest.fixture
async def client():
    transport = ASGITransport(app=app)  # type: ignore[arg-type]
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        yield c


@pytest.fixture
def docx_file(tmp_path) -> str:
    """Create a test .docx file."""
    from docx import Document

    path = str(tmp_path / "test.docx")
    doc = Document()
    doc.add_heading("Contract Title", level=1)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_paragraph("Second paragraph with details.")
    doc.save(path)
    return path


@pytest.fixture
def xlsx_file(tmp_path) -> str:
    """Create a test .xlsx file."""
    from openpyxl import Workbook

    path = str(tmp_path / "test.xlsx")
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.append(["Product", "Price", "Quantity"])
    ws.append(["Widget", "9.99", "100"])
    ws.append(["Gadget", "19.99", "50"])
    wb.save(path)
    return path


class TestDocumentExtractAPI:
    """Tests for POST /files/extract-document"""

    @pytest.mark.anyio
    async def test_extract_docx_local(self, client: AsyncClient, docx_file: str) -> None:
        """Test .docx extraction in local mode."""
        os.environ["DEPLOY_MODE"] = "tauri"
        try:
            resp = await client.post(
                "/api/v1/files/extract-document",
                json={"filePath": docx_file},
            )
        finally:
            os.environ.pop("DEPLOY_MODE", None)

        assert resp.status_code == 200
        data = resp.json()["data"]
        assert data["format"] == "docx"
        assert data["charCount"] > 0
        assert "Contract Title" in data["text"]
        assert "test paragraph" in data["text"]

    @pytest.mark.anyio
    async def test_extract_xlsx_local(self, client: AsyncClient, xlsx_file: str) -> None:
        """Test .xlsx extraction in local mode."""
        os.environ["DEPLOY_MODE"] = "tauri"
        try:
            resp = await client.post(
                "/api/v1/files/extract-document",
                json={"filePath": xlsx_file},
            )
        finally:
            os.environ.pop("DEPLOY_MODE", None)

        assert resp.status_code == 200
        data = resp.json()["data"]
        assert data["format"] == "xlsx"
        assert data["charCount"] > 0
        assert "Product" in data["text"]
        assert "Widget" in data["text"]
        assert "9.99" in data["text"]

    @pytest.mark.anyio
    async def test_missing_params(self, client: AsyncClient) -> None:
        """Test error when neither fileId nor filePath provided."""
        resp = await client.post(
            "/api/v1/files/extract-document",
            json={},
        )
        assert resp.status_code == 400

    @pytest.mark.anyio
    async def test_unsupported_format(self, client: AsyncClient, tmp_path) -> None:
        """Test error for unsupported format (.rtf)."""
        rtf_path = str(tmp_path / "document.rtf")
        with open(rtf_path, "w") as f:
            f.write("fake rtf")

        os.environ["DEPLOY_MODE"] = "tauri"
        try:
            resp = await client.post(
                "/api/v1/files/extract-document",
                json={"filePath": rtf_path},
            )
        finally:
            os.environ.pop("DEPLOY_MODE", None)

        assert resp.status_code == 400

    @pytest.mark.anyio
    async def test_file_not_found(self, client: AsyncClient) -> None:
        """Test error when file doesn't exist."""
        os.environ["DEPLOY_MODE"] = "tauri"
        try:
            resp = await client.post(
                "/api/v1/files/extract-document",
                json={"filePath": "/nonexistent/path/file.docx"},
            )
        finally:
            os.environ.pop("DEPLOY_MODE", None)

        assert resp.status_code == 404

    @pytest.mark.anyio
    async def test_empty_docx(self, client: AsyncClient, tmp_path) -> None:
        """Test extraction of an empty .docx file."""
        from docx import Document

        path = str(tmp_path / "empty.docx")
        doc = Document()
        doc.save(path)

        os.environ["DEPLOY_MODE"] = "tauri"
        try:
            resp = await client.post(
                "/api/v1/files/extract-document",
                json={"filePath": path},
            )
        finally:
            os.environ.pop("DEPLOY_MODE", None)

        assert resp.status_code == 200
        data = resp.json()["data"]
        assert data["charCount"] == 0
